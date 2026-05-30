#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
证书生成器 - Streamlit 在线版
功能：导入证书模板、导入获奖名单、生成打印文件
"""

import streamlit as st
import pandas as pd
import os
import shutil
import tempfile
import zipfile
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 页面配置
st.set_page_config(
    page_title="证书生成器",
    page_icon="📜",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============ CSS 样式 ============
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f1f1f;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .status-box {
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .status-success {
        background-color: #d4edda;
        color: #155724;
        border: 1px solid #c3e6cb;
    }
    .status-info {
        background-color: #d1ecf1;
        color: #0c5460;
        border: 1px solid #bee5eb;
    }
    .status-warning {
        background-color: #fff3cd;
        color: #856404;
        border: 1px solid #ffeeba;
    }
    .placeholder-tag {
        display: inline-block;
        background-color: #e3f2fd;
        color: #1565c0;
        padding: 2px 8px;
        border-radius: 4px;
        font-family: monospace;
        font-size: 0.9em;
        margin: 2px;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        font-weight: 500;
    }
    .data-preview {
        font-size: 0.9rem;
    }
    .template-preview {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #dee2e6;
        font-family: 'Consolas', monospace;
        font-size: 0.9rem;
        max-height: 400px;
        overflow-y: auto;
    }
</style>
""", unsafe_allow_html=True)

# ============ 初始化 Session State ============
if 'template_path' not in st.session_state:
    st.session_state.template_path = None
if 'template_bytes' not in st.session_state:
    st.session_state.template_bytes = None
if 'award_data' not in st.session_state:
    st.session_state.award_data = []
if 'template_doc' not in st.session_state:
    st.session_state.template_doc = None

# 支持的占位符
PLACEHOLDERS = ["姓名", "组别", "指导教师", "名次", "奖项"]

# ============ 核心功能函数 ============

def get_template_preview(doc):
    """获取模板预览文本"""
    if not doc:
        return "未加载模板"

    preview = []
    preview.append("=" * 50)
    preview.append("证书模板内容预览")
    preview.append("=" * 50)
    preview.append("")

    has_content = False
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            has_content = True
            preview.append(f"[段落 {i+1}] {text}")

    if not has_content:
        preview.append("【模板内容为空】")
        preview.append("请先在Word中编辑模板，添加占位符")

    preview.append("")
    preview.append("=" * 50)
    preview.append("支持的占位符:")
    for ph in PLACEHOLDERS:
        preview.append(f"  [{ph}]")
    preview.append("=" * 50)

    return "\n".join(preview)


def create_certificate_file(template_bytes, record):
    """创建单个证书文件 - 直接复制模板并替换XML中的文本"""
    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name
        tmp.write(template_bytes)

    # 解压docx文件
    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(tmp_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # 注册命名空间
        import xml.etree.ElementTree as ET
        ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')

        # 遍历所有XML文件
        for root_dir, dirs, files in os.walk(temp_dir):
            for file in files:
                if file.endswith('.xml'):
                    file_path = os.path.join(root_dir, file)
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # 替换所有占位符
                    modified = False
                    for key, value in record.items():
                        placeholder = f"[{key}]"
                        replacement = str(value) if value else ""
                        if placeholder in content:
                            content = content.replace(placeholder, replacement)
                            modified = True

                    if modified:
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(content)

        # 重新打包
        output_buffer = BytesIO()
        with zipfile.ZipFile(output_buffer, 'w') as zip_ref:
            for root_dir, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_ref.write(file_path, arcname)

        output_buffer.seek(0)
        return output_buffer.read()

    finally:
        shutil.rmtree(temp_dir)
        try:
            os.unlink(tmp_path)
        except:
            pass


def merge_docx_files(file_bytes_list):
    """通过XML合并多个docx文件，保留完整格式"""
    if not file_bytes_list:
        return None

    if len(file_bytes_list) == 1:
        return file_bytes_list[0]

    # 保存第一个文件
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(file_bytes_list[0])
        first_path = tmp.name

    # 复制第一个文件作为基础
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    shutil.copy2(first_path, output_path)

    # 解压目标文件
    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(output_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # 读取document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            root_xml = f.read()

        # 收集所有后续文档的主体内容
        bodies = []
        for file_bytes in file_bytes_list[1:]:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(file_bytes)
                file_path = tmp.name

            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)

                tmp_xml_path = os.path.join(tmp_dir, 'word', 'document.xml')
                with open(tmp_xml_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # 提取body内容
                body_match = re.search(r'<w:body[^>]*>(.*?)</w:body>', content, re.DOTALL)
                if body_match:
                    bodies.append(body_match.group(1))

            try:
                os.unlink(file_path)
            except:
                pass

        # 在第一个文档的body末尾添加分页符和其他文档内容
        insert_pos = root_xml.rfind('</w:body>')
        if insert_pos > 0 and bodies:
            page_break = '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'

            new_content = root_xml[:insert_pos]
            for body_content in bodies:
                new_content += page_break
                new_content += body_content
            new_content += root_xml[insert_pos:]

            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(new_content)

        # 重新打包
        output_buffer = BytesIO()
        with zipfile.ZipFile(output_buffer, 'w') as zip_ref:
            for root_dir, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_ref.write(file_path, arcname)

        output_buffer.seek(0)
        return output_buffer.read()

    finally:
        shutil.rmtree(temp_dir)
        try:
            os.unlink(first_path)
            os.unlink(output_path)
        except:
            pass


def create_sample_template():
    """生成示例证书模板"""
    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # 标题
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("获 奖 证 书")
    run1.bold = True
    run1.font.size = Pt(36)
    run1.font.name = "微软雅黑"

    doc.add_paragraph()

    # 内容
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("兹证明")
    run2.font.size = Pt(16)
    run2.font.name = "宋体"

    doc.add_paragraph()

    # 姓名
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("[姓名]")
    run3.bold = True
    run3.font.size = Pt(28)
    run3.font.name = "微软雅黑"

    doc.add_paragraph()

    # 组别
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("在 [组别] 比赛中")
    run4.font.size = Pt(16)
    run4.font.name = "宋体"

    doc.add_paragraph()

    # 名次
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("荣获 [名次]")
    run5.bold = True
    run5.font.size = Pt(24)
    run5.font.name = "微软雅黑"

    doc.add_paragraph()

    # 奖项
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("（[奖项]）")
    run6.font.size = Pt(18)
    run6.font.name = "宋体"

    doc.add_paragraph()
    doc.add_paragraph()

    # 指导教师
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run7 = p7.add_run("指导教师：[指导教师]")
    run7.font.size = Pt(14)
    run7.font.name = "宋体"

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 落款
    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run8 = p8.add_run("特发此证，以资鼓励")
    run8.font.size = Pt(14)
    run8.font.name = "楷体"

    doc.add_paragraph()

    p9 = doc.add_paragraph()
    p9.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run9 = p9.add_run("（盖章）")
    run9.font.size = Pt(14)
    run9.font.name = "宋体"

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============ 页面主体 ============

# 标题
st.markdown('<div class="main-header">📜 证书生成器</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">导入证书模板 + 获奖名单 → 一键生成打印文件</div>', unsafe_allow_html=True)

# 占位符说明
cols = st.columns(5)
for i, ph in enumerate(PLACEHOLDERS):
    cols[i].markdown(f'<div style="text-align:center"><span class="placeholder-tag">[{ph}]</span></div>',
                     unsafe_allow_html=True)

st.markdown("---")

# 左右分栏
left_col, right_col = st.columns([1, 1])

# ============ 左侧面板：数据与操作 ============
with left_col:
    st.subheader("📁 文件导入")

    # 导入证书模板
    st.markdown("**1. 导入证书模板 (.docx)**")
    template_file = st.file_uploader(
        "选择证书模板文件",
        type=['docx'],
        key="template_uploader",
        help="模板中可使用 [姓名]、[组别]、[指导教师]、[名次]、[奖项] 作为占位符"
    )

    if template_file is not None:
        st.session_state.template_bytes = template_file.getvalue()
        try:
            # 从内存加载文档
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(st.session_state.template_bytes)
                tmp_path = tmp.name

            st.session_state.template_doc = Document(tmp_path)
            st.session_state.template_path = template_file.name

            os.unlink(tmp_path)

            st.markdown('<div class="status-box status-success">✅ 模板导入成功: {}</div>'.format(
                template_file.name), unsafe_allow_html=True)
        except Exception as e:
            st.error(f"模板导入失败: {e}")

    # 生成示例模板按钮
    if st.button("🎨 生成示例模板", use_container_width=True):
        sample_bytes = create_sample_template()
        st.download_button(
            label="⬇️ 下载示例模板",
            data=sample_bytes,
            file_name="证书模板示例.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.success("示例模板已生成，点击上方按钮下载！")

    st.markdown("---")

    # 导入获奖名单
    st.markdown("**2. 导入获奖名单 (.xlsx / .xls)**")
    excel_file = st.file_uploader(
        "选择获奖名单Excel文件",
        type=['xlsx', 'xls'],
        key="excel_uploader",
        help="Excel需包含：姓名、组别、指导教师、名次、奖项 列"
    )

    if excel_file is not None:
        try:
            df = pd.read_excel(excel_file)
            df.columns = [str(col).strip() for col in df.columns]

            # 转换数据
            award_data = []
            for _, row in df.iterrows():
                name = str(row.get("姓名", "")).strip() if pd.notna(row.get("姓名")) else ""
                if name and name != "nan" and name != "姓名":
                    record = {
                        "姓名": name,
                        "组别": str(row.get("组别", "")).strip() if pd.notna(row.get("组别")) else "",
                        "指导教师": str(row.get("指导教师", "")).strip() if pd.notna(row.get("指导教师")) else "",
                        "名次": str(row.get("名次", "")).strip() if pd.notna(row.get("名次")) else "",
                        "奖项": str(row.get("奖项", "")).strip() if pd.notna(row.get("奖项")) else "",
                    }
                    award_data.append(record)

            st.session_state.award_data = award_data

            if award_data:
                st.markdown('<div class="status-box status-success">✅ 名单导入成功！共 {} 条记录</div>'.format(
                    len(award_data)), unsafe_allow_html=True)
            else:
                st.warning("未找到有效数据，请检查Excel文件格式")
        except Exception as e:
            st.error(f"名单导入失败: {e}")

    st.markdown("---")

    # 生成按钮
    st.markdown("**3. 生成证书文件**")

    if st.button("🚀 生成所有证书", type="primary", use_container_width=True):
        if not st.session_state.template_bytes:
            st.error("❌ 请先导入证书模板！")
        elif not st.session_state.award_data:
            st.error("❌ 请先导入获奖名单！")
        else:
            with st.spinner("正在生成证书文件，请稍候..."):
                try:
                    total = len(st.session_state.award_data)
                    file_bytes_list = []

                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    for idx, record in enumerate(st.session_state.award_data):
                        status_text.text(f"正在生成: {record['姓名']} ({idx+1}/{total})")

                        cert_bytes = create_certificate_file(
                            st.session_state.template_bytes,
                            record
                        )
                        file_bytes_list.append(cert_bytes)

                        progress_bar.progress((idx + 1) / total)

                    status_text.text("正在合并文件...")

                    # 合并所有证书
                    merged_bytes = merge_docx_files(file_bytes_list)

                    progress_bar.empty()
                    status_text.empty()

                    st.success(f"🎉 证书生成成功！共 {total} 页")

                    st.download_button(
                        label="⬇️ 下载所有证书 (.docx)",
                        data=merged_bytes,
                        file_name="所有获奖证书.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                except Exception as e:
                    st.error(f"生成失败: {e}")
                    import traceback
                    st.code(traceback.format_exc())


# ============ 右侧面板：预览 ============
with right_col:
    st.subheader("👁️ 预览")

    # 模板预览
    st.markdown("**证书模板预览**")
    if st.session_state.template_doc:
        preview_text = get_template_preview(st.session_state.template_doc)
        st.markdown(f'<div class="template-preview"><pre>{preview_text}</pre></div>',
                    unsafe_allow_html=True)
    else:
        st.info("""
        请先导入证书模板...

        提示：
        1. 模板中可使用占位符替换数据
        2. 严格保留模板的布局、字体、字号等所有格式
        3. 如果模板为空，可点击左侧【生成示例模板】创建示例
        """)

    st.markdown("---")

    # 获奖名单预览
    st.markdown("**获奖名单预览**")
    if st.session_state.award_data:
        df_display = pd.DataFrame(st.session_state.award_data)
        st.dataframe(
            df_display,
            use_container_width=True,
            hide_index=True,
            height=300
        )
        st.markdown(f"**共 {len(st.session_state.award_data)} 条记录**")
    else:
        st.info("暂无获奖名单数据，请先导入Excel文件")


# ============ 底部信息 ============
st.markdown("---")
footer_cols = st.columns(2)
with footer_cols[0]:
    st.markdown("📌 **支持占位符**: " + " ".join([f"`[{ph}]`" for ph in PLACEHOLDERS]))
with footer_cols[1]:
    st.markdown("📊 **获奖记录**: {} 条".format(len(st.session_state.award_data)))